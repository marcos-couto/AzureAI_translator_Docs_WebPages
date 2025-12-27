import requests
from docx import Document
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
subscription_key = os.getenv("AZURE_TRANSLATOR_KEY")
endpoint = os.getenv("AZURE_TRANSLATOR_ENDPOINT")
location = os.getenv("AZURE_TRANSLATOR_REGION")
target_language = "pt-br"

def translator_document(text, target_language):
    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(os.urandom(16))
    }

    body = [{ "text": text }]
    params = { "from": "en", "to": target_language }

    response = requests.post(endpoint, params=params, headers=headers, json=body)
    result = response.json()
    return result[0]["translations"][0]["text"]

def translate_document(path):
    document = Document(path)
    full_text = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            translated_text = translator_document(paragraph.text, target_language)
            full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)

    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_doc.save(path_translated)
    return path_translated

if __name__ == "__main__":
    input_file = "Marcos_Couto_bio.docx"
    output_file = translate_document(input_file)
    print(f"Translated document saved as: {output_file}")
